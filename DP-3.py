import numpy as np
import warnings

import pandas as pd
from sklearn.metrics import classification_report
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsClassifier
from sklearn.svm import SVC
from sklearn.naive_bayes import GaussianNB
from sklearn.metrics import confusion_matrix, ConfusionMatrixDisplay
from sklearn import tree
from sklearn.metrics import accuracy_score
from imblearn.over_sampling import SMOTE
from sklearn.model_selection import KFold, cross_val_score

import docx2txt
import spacy
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

warnings.filterwarnings("ignore")

# Pre-Processing

df = pd.read_csv("Company.csv")
df = df.iloc[:, 1:]

df.head(100)

print(df.isnull().sum().sum())
df.dropna(inplace=True)

dupp = df[df.duplicated(keep=False)].shape[0]

df.head(100)

Ros = df['Risk level']

s = df['Risk level'].value_counts().rename_axis('Risk level').rename('count')  # if x- and y-labels are important

X = df.iloc[:, :-1].values
# Now let's tell the dataframe which column we want for the target/labels.  
y = df['Risk level']

df.corr()

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=27)

SVC_model = SVC()
# KNN model requires you to specify n_neighbors,
# the number of points the classifier will look at to determine what class a new point belongs to
KNN_model = KNeighborsClassifier(n_neighbors=5)

classifier = LogisticRegression()
classifier.fit(X_train, y_train)

KNN_model.fit(X_train, y_train)

SVC_model.fit(X_train, y_train)
KNN_model.fit(X_train, y_train)

SVC_prediction = SVC_model.predict(X_test)
KNN_prediction = KNN_model.predict(X_test)
logi = classifier.predict(X_test)

# SVM Method
print("----------SVM confusion matrix-----")
cmd = confusion_matrix(y_test, SVC_prediction)

print(cmd)
cmd = ConfusionMatrixDisplay(cmd, display_labels=['0', '1', '2', '3', '4', '5'])
print("SVM Accuracy:", accuracy_score(y_test, SVC_prediction))
target_names = ['class 0', 'class 1', 'class 2', 'class 3', 'class 4', 'class 5']
cmd.plot()
print(classification_report(y_test, SVC_prediction, target_names=target_names))

y_prob = KNN_model.predict_proba(X_test)
y_prob_logi = KNN_model.predict_proba(X_test)
print(accuracy_score(y_test, KNN_prediction))
print(accuracy_score(y_test, logi))

print("The risk label is :", KNN_prediction[3])

print("Accuracy:", accuracy_score(y_test, KNN_prediction))

gnb = GaussianNB()

# Train the classifier:
model = gnb.fit(X_train, y_train)
# Make predictions with the classifier:
predictive_labels = gnb.predict(X_test)

clf = tree.DecisionTreeClassifier()
clf = clf.fit(X_train, y_train)

clf_predict = clf.predict(X_test)
print("Accuracy:", accuracy_score(y_test, clf_predict))

oversample_ = SMOTE()

X_over, y_over = oversample_.fit_resample(X, y)
X_train_o, X_test_o, y_train_o, y_test_o = train_test_split(X_over, y_over, test_size=0.3, random_state=27)

SVC_model = SVC()
# KNN model requires you to specify n_neighbors,
# the number of points the classifier will look at to determine what class a new point belongs to
KNN_model = KNeighborsClassifier(n_neighbors=5)
Tree_model = tree.DecisionTreeClassifier()

# KNN model requires you to specify n_neighbors,
# the number of points the classifier will look at to determine what class a new point belongs to
KNN_model_over = KNeighborsClassifier(n_neighbors=5)

KNN_model_over.fit(X_train_o, y_train_o)
SVC_model.fit(X_train_o, y_train_o)
SVC_prediction = SVC_model.predict(X_test_o)
Tree_model.fit(X_train_o, y_train_o)
tree_prediction = Tree_model.predict(X_test_o)

KNN_predictions = KNN_model_over.predict(X_test_o)
print(accuracy_score(y_test_o, KNN_predictions))
print(accuracy_score(y_test_o, SVC_prediction))
print(accuracy_score(y_test_o, tree_prediction))

svc_over = SVC_model.predict(X_test)
print(accuracy_score(y_test, svc_over))

print("----------KNN confusion matrix-----")
cmd_ = confusion_matrix(y_test_o, KNN_predictions)

print(cmd)
cmd = ConfusionMatrixDisplay(cmd_, display_labels=['0', '1', '2', '3', '4', '5'])
print("KNN Accuracy:", accuracy_score(y_test_o, KNN_predictions))
target_names = ['class 0', 'class 1', 'class 2', 'class 3', 'class 4', 'class 5']
cmd.plot()
print(classification_report(y_test_o, KNN_predictions, target_names=target_names))

print("----------KNN confusion matrix-----")
cmd_ = confusion_matrix(y_test_o, SVC_prediction)

print(cmd)
cmd = ConfusionMatrixDisplay(cmd_, display_labels=['0', '1', '2', '3', '4', '5'])
print("SVM Accuracy:", accuracy_score(y_test_o, SVC_prediction))
target_names = ['class 0', 'class 1', 'class 2', 'class 3', 'class 4', 'class 5']
cmd.plot()
print(classification_report(y_test_o, SVC_prediction, target_names=target_names))

print("----------KNN confusion matrix-----")
cmd_ = confusion_matrix(y_test_o, tree_prediction)

print(cmd)
cmd = ConfusionMatrixDisplay(cmd_, display_labels=['0', '1', '2', '3', '4', '5'])
print("KNN Accuracy:", accuracy_score(y_test_o, tree_prediction))
target_names = ['class 0', 'class 1', 'class 2', 'class 3', 'class 4', 'class 5']
cmd.plot()
print(classification_report(y_test_o, tree_prediction, target_names=target_names))

# K-Fold Cross Validation

kfold = KFold(n_splits=5, shuffle=True, random_state=42)

# Perform cross-validation
scores = cross_val_score(KNN_model, X_over, y_over, cv=kfold)
scores_svm = cross_val_score(SVC_model, X_over, y_over, cv=kfold)
scores_clf = cross_val_score(Tree_model, X_over, y_over, cv=kfold)
# Print the cross-validation scores
print('KNN-Cross-validation scores: {}'.format(scores))
print('SVM scores: {}'.format(scores_svm))
print('Decision Tree: {}'.format(scores_clf))

phone_number_regex = r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
address_regex = "^(\\d{1,}) [a-zA-Z0-9\\s]+(\\,)? [a-zA-Z]+(\\,)? [A-Z]{2} [0-9]{5,6}$"


####################################################################
print("\n\nResume Scraper\n\n")


def txt_input(docx):
    txt = docx2txt.process(docx)
    if txt:
        return txt.replace('\t', ' ')
    return None


def find_text(resume, headers):
    # Define a list of possible section footers
    footers = ['Education', 'Skills', 'Projects', 'Experiences', 'Certificate']
    # Loop over the possible headers and find the first occurrence in the text
    for header in headers:
        match = re.search(rf"\b{header}\b", resume)
        if match:
            for footer in footers:
                if footer == header:
                    continue
                f_match = re.search(rf"\b{footer}\b", resume[match.start():])
                if f_match:
                    return match.start() + len(header), f_match.start() + match.start()
            return match.start() + len(header), len(resume)

    # If no header is found, return None
    return None


def generatorCV(risk, name, address, phone, email, education, skill, project, experience, certificate):
    doc = Document()
    doc.add_heading(name, level=1).style.font.size = Pt(20)

    info = doc.add_paragraph()
    info.add_run(address + '\n')
    info.add_run(phone + '\n')
    info.add_run(email + '\n')

    education = education.replace("\n\n", "\n").strip()
    doc.add_heading("EDUCATION", level=1)
    educationCV = doc.add_paragraph()
    educationCV.add_run(education)

    skill = skill.replace("\n\n", "\n").strip()
    doc.add_heading("SKILLS", level=1)
    skillCV = doc.add_paragraph()
    skillCV.add_run(skill)

    project = project.replace("\n\n", "\n").strip()
    doc.add_heading("PROJECTS", level=1)
    projectCV = doc.add_paragraph()
    projectCV.add_run(project)

    experience = experience.replace("\n\n", "\n").strip()
    doc.add_heading("EXPERIENCES", level=1)
    experienceCV = doc.add_paragraph()
    experienceCV.add_run(experience)

    certificate = certificate.replace("\n\n", "\n").strip()
    doc.add_heading("CERTIFICATES", level=1)
    certificateCV = doc.add_paragraph()
    certificateCV.add_run(certificate)

    if risk != 0:
        note = doc.add_paragraph("\n\n\nThis resume is being protected by applicant's preferences")
        note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        note.style.font.size = Pt(9)
        note.bold = True
    doc.save('risk_' + str(risk) + '_resume.docx')


def main(risk):
    resume_text = txt_input('AnTranResume.docx')

    nlp = spacy.load('en_core_web_sm')
    doc = nlp(resume_text)

    address = education = project = experience = skill = certificate = ""

    # This is to locate address
    lines = resume_text.splitlines()
    for line in lines:
        newline = line.splitlines()
        # Remove 'Address: '
        for i in range(len(newline)):
            newline[i] = newline[i].split(' ', 1)[-1]

        # join lines back into a single string
        new_text = '\n'.join(newline)
        matches = re.match(address_regex, new_text)
        if matches:
            # print(matches.group(0))
            address = matches.group(0)

    # This is to locate phone number
    phone = re.findall(phone_number_regex, resume_text)

    # Locate name and email address
    name = doc.ents[0] if doc.ents else ''
    name = name.text.rstrip()
    email = [token.text for token in doc if token.like_email][0] if any(token.like_email for token in doc) else ''

    # Find Education field
    education_start, education_end = find_text(resume_text, ['Education'])
    if education_start is not None:
        education = resume_text[education_start:education_end]
        education = education.rstrip()
    else:
        print("Experience section not found")

    # Find Skill Section
    skill_start, skill_end = find_text(resume_text, ['Skills'])
    if skill_start is not None:
        skill = resume_text[skill_start:skill_end]
        skill = skill.rstrip()
    else:
        print("Skill section not found")

    # Find Project Section
    project_start, project_end = find_text(resume_text, ['Projects'])
    if project_start is not None:
        project = resume_text[project_start:project_end]
        project = project.rstrip()
    else:
        print("Project section not found")

    # Find Experience Section
    experience_start, experience_end = find_text(resume_text, ['Experiences'])
    if experience_start is not None:
        experience = resume_text[experience_start:experience_end]
        experience = experience.rstrip()
    else:
        print("Experience section not found")

    # Find Certificate Section
    certificate_start, certificate_end = find_text(resume_text, ['Certificate'])
    if certificate_start is not None:
        certificate = resume_text[certificate_start:certificate_end]
        certificate = certificate.rstrip()
    else:
        print("Certificate section not found")

    # Masking information
    at_index = email.index("@")
    email_covered = "*" * at_index + email[at_index:]

    phone_covered = phone[0][:-4] + "****"

    comma_index = address.index(",")
    address_covered = "*" * comma_index + address[comma_index:]

    education_covered = education.splitlines()
    for i in range(len(education_covered)):
        for school in ["University", "College"]:
            if school in education_covered[i]:
                education_covered[i] = school + " of " + "".join(["*" for _ in education_covered[i]])
    education_covered = "\n".join(education_covered)

    experience_covered = experience.splitlines()
    for i in range(len(experience_covered)):
        if "-" in experience_covered[i]:
            dash_index = experience_covered[i].index("-")
            experience_covered[i] = experience_covered[i][:dash_index] + "*" * (len(experience_covered[i]) - dash_index)
        else:
            continue
    experience_covered = "\n".join(experience_covered)

    project_covered = project.splitlines()
    for i in range(len(project_covered)):
        if "-" in project_covered[i]:
            dash_index = project_covered[i].index("-")
            project_covered[i] = project_covered[i][:dash_index] + "*" * (len(project_covered[i]) - dash_index)
        else:
            continue
    project_covered = "\n".join(project_covered)

    # Generate CV
    risk = int(risk)
    if risk == 0:
        generatorCV(risk, name, address, phone[0], email,
                    education, skill, project, experience, certificate)
    elif risk == 1:
        generatorCV(risk, name, address_covered, phone[0], email,
                    education, skill, project, experience, certificate)
    elif risk == 2:
        generatorCV(risk, name, address_covered, phone_covered, email,
                    education, skill, project, experience, certificate)
    elif risk == 3:
        generatorCV(risk, name, address_covered, phone_covered, email_covered,
                    education, skill, project, experience, certificate)
    elif risk == 4:
        generatorCV(risk, name, address_covered, phone_covered, email_covered,
                    education_covered, skill, project, experience, certificate)
    elif risk == 5:
        generatorCV(risk, name, address_covered, phone_covered, email_covered,
                    education_covered, skill, project_covered, experience_covered, certificate)


A = np.array([[0, 0, 0, 0, 0, 0, 1, 0, 1, 0, 0, 1]])
# using inputs to predict the output
prediction = KNN_model_over.predict(A)
print("Company A risk level: {}".format(prediction[0]))

riskLevel = prediction[0]
main(riskLevel)

B = np.array([[0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]])
# using inputs to predict the output
prediction = KNN_model_over.predict(B)
print("Company B risk level: {}".format(prediction[0]))

riskLevel = prediction[0]
main(riskLevel)
